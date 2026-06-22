from datetime import date
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TEMPLATE = Path(r"C:\Users\Aleja\Desktop\Plantilla proyecto Proyecto final (3).docx")
OUTPUT = Path(r"C:\Users\Aleja\Desktop\Documentacion_UniMarket_Final.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("• ")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    p.add_run(text)
    return p


def add_number(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{number}. ")
    run.font.size = Pt(11)
    p.add_run(text)
    return p


def get_numbering_ids(doc):
    if hasattr(doc, "_unimarket_numbering_ids"):
        return doc._unimarket_numbering_ids

    numbering = doc.part.numbering_part.element
    abstract_ids = []
    num_ids = []
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        value = abstract_num.get(qn("w:abstractNumId"))
        if value is not None and value.isdigit():
            abstract_ids.append(int(value))
    for num in numbering.findall(qn("w:num")):
        value = num.get(qn("w:numId"))
        if value is not None and value.isdigit():
            num_ids.append(int(value))

    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    bullet_abstract = add_numbering_definition(
        numbering,
        abstract_id=next_abstract,
        fmt="bullet",
        text="•",
    )
    bullet_num = add_numbering_instance(numbering, num_id=next_num, abstract_id=bullet_abstract)

    number_abstract = add_numbering_definition(
        numbering,
        abstract_id=next_abstract + 1,
        fmt="decimal",
        text="%1.",
    )
    number_num = add_numbering_instance(numbering, num_id=next_num + 1, abstract_id=number_abstract)

    doc._unimarket_numbering_ids = (bullet_num, number_num)
    return doc._unimarket_numbering_ids


def add_numbering_definition(numbering, abstract_id, fmt, text):
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)

    if fmt == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(r_fonts)
        lvl.append(r_pr)

    abstract_num.append(lvl)
    numbering.append(abstract_num)
    return abstract_id


def add_numbering_instance(numbering, num_id, abstract_id):
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")

    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_page_break(doc):
    doc.add_page_break()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "E8EEF5")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for i, value in enumerate(row_data):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MARKETPLACE UNIVERSITARIO\nUNIMARKET")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(46, 116, 181)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documentacion del sistema movil")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(89, 89, 89)

    for _ in range(5):
        doc.add_paragraph()

    lines = [
        "Presentado por: Mendoza, Alejandra",
        "Carrera: Ingenieria en Sistemas",
        "Asignatura: Proyecto final",
        "Ciudad: Guayaquil",
        "Fecha: 18 de junio de 2026",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)


def build_document():
    shutil.copyfile(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document_body(doc)
    configure_styles(doc)

    add_cover(doc)
    add_page_break(doc)

    add_heading(doc, "Resumen", 1)
    add_paragraph(
        doc,
        "El presente documento describe el desarrollo de UniMarket, una aplicacion movil Android orientada "
        "a facilitar la compra, venta, publicacion y comunicacion entre estudiantes universitarios. El sistema "
        "fue implementado en Kotlin con Material Design 3, navegacion mediante una sola Activity principal y "
        "varios Fragments para representar las interfaces de login, registro, inicio, detalle de producto, "
        "publicacion, favoritos, chat y perfil. En esta primera version se utiliza un repositorio local en memoria "
        "para simular usuarios, productos, favoritos y mensajes. Adicionalmente, se preparo un script de base de "
        "datos PostgreSQL con tablas para usuarios, categorias, productos, multimedia, favoritos, conversaciones "
        "y mensajes, con el objetivo de conectar posteriormente la aplicacion mediante una API REST segura.",
    )

    add_heading(doc, "Indice de contenidos", 1)
    for item in [
        "1. Introduccion",
        "1.1 Justificacion",
        "1.2 Planteamiento del trabajo",
        "1.3 Estructura de la memoria",
        "2. Contexto y estado del arte",
        "3. Objetivos concretos y metodologia de trabajo",
        "3.4 Metodologia Scrum adaptada",
        "4. Desarrollo especifico de la contribucion",
        "5. Conclusiones y trabajo futuro",
        "6. Bibliografia",
        "Anexos",
    ]:
        add_paragraph(doc, item)

    add_heading(doc, "Indice de tablas", 1)
    for item in [
        "Tabla 1. Sprints de la metodologia Scrum adaptada",
        "Tabla 2. Tecnologias utilizadas",
        "Tabla 3. Requisitos funcionales principales",
        "Tabla 4. Interfaces del sistema",
        "Tabla 5. Modelo de datos principal",
        "Tabla 6. Eventos principales de la aplicacion",
    ]:
        add_paragraph(doc, item)

    add_heading(doc, "Indice de figuras", 1)
    add_paragraph(doc, "Figura 1. Arquitectura general propuesta para UniMarket")

    add_heading(doc, "1. Introduccion", 1)
    add_paragraph(
        doc,
        "UniMarket es una aplicacion movil de tipo marketplace universitario. Su finalidad es permitir que los "
        "estudiantes publiquen productos o servicios academicos, consulten publicaciones disponibles, filtren por "
        "categoria, guarden favoritos y contacten a vendedores mediante un chat interno. El desarrollo se realizo "
        "en Android Studio usando Kotlin, XML layouts, ViewBinding, Material Design 3 y Navigation Component.",
    )
    add_paragraph(
        doc,
        "La aplicacion se encuentra en una primera etapa funcional sin backend conectado. Para esta fase se implemento "
        "un repositorio local llamado MarketplaceRepository que almacena datos simulados en memoria. Esto permite "
        "probar la navegacion, las interfaces y los eventos antes de conectar la persistencia real con PostgreSQL.",
    )

    add_heading(doc, "1.1 Justificacion", 2)
    add_paragraph(
        doc,
        "En el entorno universitario es comun que los estudiantes necesiten comprar o vender libros, calculadoras, "
        "materiales de laboratorio, accesorios, servicios de tutoria u otros recursos academicos. Sin embargo, esta "
        "actividad suele realizarse mediante grupos informales de mensajeria o redes sociales, lo que dificulta la "
        "busqueda por categoria, el seguimiento de publicaciones y la comunicacion ordenada entre comprador y vendedor.",
    )
    add_paragraph(
        doc,
        "UniMarket responde a esta necesidad mediante una solucion movil centrada en la comunidad universitaria. El "
        "sistema organiza las publicaciones por categorias, permite gestionar favoritos y ofrece una base para integrar "
        "mensajeria, autenticacion y almacenamiento permanente.",
    )

    add_heading(doc, "1.2 Planteamiento del trabajo", 2)
    add_paragraph(
        doc,
        "El trabajo propone el desarrollo de una aplicacion Android que funcione como prototipo operativo de un "
        "marketplace universitario. La solucion actual incluye las pantallas principales, flujo de navegacion, "
        "manejo de eventos y datos locales de prueba. Para una segunda fase se plantea la conexion de la aplicacion "
        "con una API REST y una base de datos PostgreSQL.",
    )

    add_heading(doc, "1.3 Estructura de la memoria", 2)
    add_paragraph(
        doc,
        "La memoria se organiza en seis capitulos. El primero presenta la introduccion y justificacion. El segundo "
        "describe el contexto tecnologico. El tercero establece objetivos y metodologia. El cuarto detalla el "
        "desarrollo de la aplicacion. El quinto presenta conclusiones y trabajo futuro. Finalmente, se incluye una "
        "bibliografia y anexos tecnicos con los archivos principales del proyecto.",
    )

    add_heading(doc, "2. Contexto y estado del arte", 1)
    add_paragraph(
        doc,
        "Las aplicaciones de marketplace permiten publicar productos, consultar catalogos, aplicar filtros, guardar "
        "elementos de interes y comunicar compradores con vendedores. En el caso universitario, estas funciones deben "
        "adaptarse a un contexto mas cerrado, donde los usuarios pertenecen a una institucion y los productos estan "
        "relacionados con necesidades academicas o de vida estudiantil.",
    )
    add_paragraph(
        doc,
        "Para el desarrollo movil se utilizo Android Studio con Kotlin. La interfaz se construyo mediante archivos XML "
        "y componentes Material Design 3. La navegacion se implemento con Navigation Component, usando una sola Activity "
        "principal y diferentes Fragments. Esta organizacion facilita separar las pantallas sin crear una Activity para "
        "cada vista.",
    )

    add_caption(doc, "Figura 1. Arquitectura general propuesta para UniMarket")
    add_paragraph(
        doc,
        "Android App (MainActivity + Fragments) -> API REST futura -> PostgreSQL",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "3. Objetivos concretos y metodologia de trabajo", 1)
    add_heading(doc, "3.1 Objetivo general", 2)
    add_paragraph(
        doc,
        "Desarrollar una aplicacion movil Android para un marketplace universitario que permita registrar usuarios, "
        "gestionar publicaciones, buscar productos por categoria, administrar favoritos y facilitar la comunicacion "
        "entre compradores y vendedores.",
    )

    add_heading(doc, "3.2 Objetivos especificos", 2)
    for item in [
        "Permitir el registro e inicio de sesion de usuarios.",
        "Gestionar publicaciones de productos mediante creacion, edicion, eliminacion y marcado como vendido.",
        "Facilitar la busqueda y filtrado de productos por categoria.",
        "Permitir la comunicacion entre comprador y vendedor mediante una pantalla de chat.",
        "Gestionar una lista de productos favoritos por usuario.",
        "Preparar la estructura de base de datos PostgreSQL para una futura integracion con API REST.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3 Metodologia del trabajo", 2)
    for index, item in enumerate([
        "Analisis de requisitos a partir de las funcionalidades solicitadas para el marketplace.",
        "Diseno de interfaces usando XML y componentes Material Design 3.",
        "Implementacion de navegacion con una Activity principal y Fragments.",
        "Creacion de un repositorio local en memoria para simular datos mientras no existe backend.",
        "Elaboracion del script PostgreSQL con las tablas necesarias para la persistencia futura.",
        "Verificacion mediante compilacion del APK debug y ejecucion de pruebas unitarias locales.",
    ], start=1):
        add_number(doc, index, item)

    add_heading(doc, "3.4 Metodologia Scrum adaptada", 2)
    add_paragraph(
        doc,
        "Para el desarrollo de UniMarket se aplico una metodologia Scrum adaptada al alcance del proyecto academico. "
        "El trabajo se dividio en sprints cortos, donde cada sprint entrego una parte verificable del sistema. Esta "
        "forma de trabajo permitio avanzar por etapas, revisar resultados y ajustar la aplicacion conforme se "
        "implementaban nuevas funcionalidades.",
    )
    add_paragraph(
        doc,
        "Los roles Scrum fueron adaptados de la siguiente manera: el Product Owner fue representado por la estudiante, "
        "quien definio las necesidades del sistema; el Scrum Master se asumio como rol de organizacion y seguimiento "
        "del avance; y el equipo de desarrollo se encargo de implementar la aplicacion movil, preparar la base de datos "
        "y elaborar la documentacion.",
    )
    add_table(
        doc,
        ["Sprint", "Objetivo", "Entregable"],
        [
            ["Sprint 1", "Analizar requisitos del marketplace universitario.", "Lista de funcionalidades: login, registro, publicaciones, favoritos, chat y perfil."],
            ["Sprint 2", "Disenar las interfaces principales.", "Layouts XML con Material Design 3 e identificadores en espanol."],
            ["Sprint 3", "Implementar la funcionalidad movil.", "MainActivity, Fragments, navegacion y eventos con setOnClickListener."],
            ["Sprint 4", "Preparar la persistencia futura.", "Script PostgreSQL con tablas, relaciones, indices y datos de prueba."],
            ["Sprint 5", "Verificar y documentar el sistema.", "APK debug compilado, pruebas locales ejecutadas y documento final en Word."],
        ],
        [1400, 3600, 4360],
    )
    add_caption(doc, "Tabla 1. Sprints de la metodologia Scrum adaptada")

    add_page_break(doc)
    add_heading(doc, "4. Desarrollo especifico de la contribucion", 1)
    add_heading(doc, "4.1 Tecnologias utilizadas", 2)
    add_table(
        doc,
        ["Capa", "Tecnologia", "Uso en el sistema"],
        [
            ["Frontend movil", "Android Studio + Kotlin", "Implementacion principal de la aplicacion Android."],
            ["Interfaz", "XML layouts + Material Design 3", "Diseno visual de formularios, botones, tarjetas, chips y pantallas."],
            ["Navegacion", "Navigation Component", "Cambio entre Fragments desde una sola Activity principal."],
            ["Binding", "ViewBinding", "Acceso seguro a vistas declaradas en XML."],
            ["Backend futuro", "API REST", "Comunicacion futura entre Android y la base de datos."],
            ["Base de datos", "PostgreSQL", "Persistencia futura de usuarios, productos, multimedia, favoritos y mensajes."],
        ],
        [1800, 2500, 5060],
    )
    add_caption(doc, "Tabla 2. Tecnologias utilizadas")

    add_heading(doc, "4.2 Requisitos funcionales principales", 2)
    add_table(
        doc,
        ["Modulo", "Requisito", "Estado actual"],
        [
            ["Usuarios", "Registro, inicio de sesion, perfil y cierre de sesion.", "Implementado visualmente con datos locales."],
            ["Publicaciones", "Crear, editar, eliminar, consultar y marcar como vendido.", "Implementado con repositorio en memoria."],
            ["Categorias", "Libros, Electronicos, Laboratorio, Tutorias y Otros.", "Implementado con chips y selector Spinner."],
            ["Multimedia", "Agregar imagenes y video demostrativo.", "Simulado mediante selectores de contenido."],
            ["Favoritos", "Agregar, eliminar y consultar favoritos.", "Implementado con conjunto local de favoritos."],
            ["Mensajeria", "Contactar vendedor, enviar y recibir mensajes.", "Implementado como chat simulado."],
        ],
        [1600, 4560, 3200],
    )
    add_caption(doc, "Tabla 3. Requisitos funcionales principales")

    add_heading(doc, "4.3 Diseno de la aplicacion Android", 2)
    add_paragraph(
        doc,
        "El sistema utiliza una sola Activity llamada MainActivity. Esta Activity funciona como contenedor principal "
        "de navegacion y aloja el NavHostFragment. Las pantallas no fueron creadas como Activities independientes, "
        "sino como Fragments. Esta decision permite un flujo mas ordenado y facilita centralizar la navegacion.",
    )
    add_paragraph(
        doc,
        "Fragments implementados: LoginFragment, RegisterFragment, HomeFragment, ProductDetailFragment, "
        "PublishProductFragment, FavoritesFragment, ChatFragment y ProfileFragment.",
    )

    add_page_break(doc)
    add_heading(doc, "4.4 Interfaces del sistema", 2)
    add_table(
        doc,
        ["Pantalla", "Archivo XML", "Funcion principal"],
        [
            ["Login", "fragmento_inicio_sesion.xml", "Autenticar al usuario mediante correo y contrasena."],
            ["Registro", "fragmento_registro.xml", "Registrar nombre, apellido, correo, contrasena, carrera y telefono."],
            ["Inicio", "fragmento_inicio.xml", "Mostrar publicaciones, busqueda, categorias, favoritos, perfil y publicar producto."],
            ["Detalle producto", "fragmento_detalle_producto.xml", "Mostrar informacion completa del producto y contactar vendedor."],
            ["Publicar producto", "fragmento_publicar_producto.xml", "Crear o editar publicaciones con titulo, descripcion, precio, categoria y multimedia."],
            ["Favoritos", "fragmento_favoritos.xml", "Listar publicaciones guardadas y permitir eliminarlas."],
            ["Chat", "fragmento_chat.xml", "Enviar y visualizar mensajes entre comprador y vendedor."],
            ["Perfil", "fragmento_perfil.xml", "Mostrar datos personales, publicaciones propias, ventas y favoritos."],
        ],
        [1850, 2850, 4660],
    )
    add_caption(doc, "Tabla 4. Interfaces del sistema")

    add_heading(doc, "4.5 Modelo de datos local", 2)
    add_paragraph(
        doc,
        "En la version actual se usa el archivo MarketplaceRepository.kt como fuente de datos temporal. En este archivo "
        "se definieron constructores Kotlin mediante data class y enum class con nombres en espanol para facilitar la "
        "defensa del proyecto.",
    )
    add_table(
        doc,
        ["Modelo", "Campos principales", "Descripcion"],
        [
            ["Usuario", "nombre, apellido, correo, carrera, telefono", "Representa al estudiante registrado en la aplicacion."],
            ["Producto", "id, titulo, descripcion, precio, categoria, estado, vendedor, multimedia", "Representa una publicacion dentro del marketplace."],
            ["CategoriaProducto", "LIBROS, ELECTRONICOS, LABORATORIO, TUTORIAS, OTROS", "Enumera las categorias disponibles."],
            ["MensajeChat", "remitente, contenido, enviadoPorUsuarioActual", "Representa los mensajes enviados en una conversacion."],
        ],
        [1900, 3860, 3600],
    )
    add_caption(doc, "Tabla 5. Modelo de datos principal")

    add_page_break(doc)
    add_heading(doc, "4.6 Eventos y navegacion", 2)
    add_paragraph(
        doc,
        "Los eventos de la aplicacion se generan mediante listeners de Android, principalmente setOnClickListener. "
        "Un boton no crea una Activity; el boton dispara un evento y el codigo dentro del listener ejecuta una accion, "
        "por ejemplo navegar a otro Fragment, guardar un producto, eliminar un favorito o enviar un mensaje.",
    )
    add_table(
        doc,
        ["Evento", "Accion ejecutada", "Pantalla relacionada"],
        [
            ["botonIniciarSesion", "Valida campos y navega al inicio.", "Login"],
            ["botonRegistrarse", "Navega a la pantalla de registro.", "Login"],
            ["botonPublicarProducto", "Navega al formulario de publicacion.", "Inicio"],
            ["botonFavorito", "Agrega o elimina una publicacion de favoritos.", "Inicio y detalle"],
            ["botonContactarVendedor", "Abre la pantalla de chat del producto.", "Detalle producto"],
            ["botonPublicar", "Crea o actualiza una publicacion.", "Publicar producto"],
            ["botonEnviar", "Agrega el mensaje al historial del chat.", "Chat"],
            ["botonCerrarSesion", "Regresa a la pantalla de inicio de sesion.", "Perfil"],
        ],
        [2300, 4160, 2900],
    )
    add_caption(doc, "Tabla 6. Eventos principales de la aplicacion")

    add_heading(doc, "4.7 Base de datos PostgreSQL", 2)
    add_paragraph(
        doc,
        "Para la persistencia futura se preparo el archivo database/unimarket_postgresql.sql. El script crea las tablas "
        "users, categories, products, product_media, favorites, conversations y messages. Tambien incluye indices, "
        "triggers de actualizacion y datos de prueba. La aplicacion Android no debe conectarse directamente a PostgreSQL; "
        "lo recomendable es crear una API REST intermedia para proteger credenciales y controlar la logica de negocio.",
    )

    add_heading(doc, "4.8 Pruebas realizadas", 2)
    add_paragraph(
        doc,
        "La aplicacion fue verificada mediante la compilacion del APK debug y la ejecucion de pruebas unitarias locales. "
        "Los comandos ejecutados fueron:",
    )
    add_bullet(doc, ".\\gradlew.bat assembleDebug")
    add_bullet(doc, ".\\gradlew.bat testDebugUnitTest")
    add_paragraph(
        doc,
        "Ambos procesos finalizaron correctamente, por lo que el proyecto compila y genera el APK debug sin errores.",
    )

    add_page_break(doc)
    add_heading(doc, "5. Conclusiones y trabajo futuro", 1)
    add_heading(doc, "5.1 Conclusiones", 2)
    add_paragraph(
        doc,
        "El desarrollo de UniMarket permitio construir una base funcional para un marketplace universitario. La aplicacion "
        "ya cuenta con pantallas principales, navegacion completa, eventos de botones, datos simulados y una estructura "
        "preparada para persistencia real. La organizacion mediante una sola Activity y varios Fragments facilita el "
        "mantenimiento del proyecto y permite explicar con claridad la arquitectura durante la defensa.",
    )

    add_heading(doc, "5.2 Lineas de trabajo futuro", 2)
    for item in [
        "Implementar una API REST con Node.js o Spring Boot.",
        "Conectar la API REST con la base de datos PostgreSQL.",
        "Reemplazar MarketplaceRepository por servicios remotos.",
        "Agregar autenticacion real con contrasenas cifradas y tokens.",
        "Subir imagenes y videos a almacenamiento seguro.",
        "Implementar chat persistente e historial real de conversaciones.",
        "Agregar validacion de correo institucional.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Bibliografia", 1)
    for item in [
        "Android Developers. Documentacion oficial de Android y Kotlin para el desarrollo de aplicaciones moviles.",
        "Google Material Design. Guia de componentes y principios de Material Design 3.",
        "PostgreSQL Global Development Group. Documentacion oficial de PostgreSQL.",
        "Documentacion interna del proyecto UniMarket desarrollada en Android Studio.",
    ]:
        add_paragraph(doc, item)

    add_heading(doc, "Anexos", 1)
    add_heading(doc, "Anexo A. Archivos principales del proyecto", 2)
    add_paragraph(
        doc,
        "Los archivos principales del proyecto son MainActivity.kt, MarketplaceRepository.kt, los fragments Kotlin "
        "de login, registro, inicio, detalle, publicacion, favoritos, chat y perfil, ademas de los XML "
        "fragmento_inicio_sesion.xml, fragmento_registro.xml, fragmento_inicio.xml, fragmento_detalle_producto.xml, "
        "fragmento_publicar_producto.xml, fragmento_favoritos.xml, fragmento_chat.xml y fragmento_perfil.xml.",
    )
    add_paragraph(
        doc,
        "La navegacion se encuentra en grafo_navegacion.xml y la propuesta de persistencia se encuentra en "
        "database/unimarket_postgresql.sql.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
